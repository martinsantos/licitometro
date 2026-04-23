"""
UMKnowledgeService — Base de conocimiento vectorial de Ultima Milla.

Ingesta documentos internos (XLS markup, DOCX propuestas, PDF/imágenes actas)
y los hace disponibles para el Cotizador via búsqueda semántica por cosine similarity.

Colección MongoDB: um_knowledge
- chunk por documento (~400 tokens, overlap 80 chars)
- vector 384-dim (paraphrase-multilingual-MiniLM-L12-v2, mismo modelo que embedding_service)
- búsqueda in-memory con numpy (misma estrategia que embedding_service)
"""
import asyncio
import logging
import os
import re
import uuid
from datetime import datetime, timezone
from typing import Optional

from motor.motor_asyncio import AsyncIOMotorDatabase

logger = logging.getLogger("um_knowledge_service")

CHUNK_SIZE = 1800   # ~400 tokens
CHUNK_OVERLAP = 120

_USE_EMBEDDINGS = os.getenv("USE_EMBEDDINGS", "false").lower() == "true"

TIPOS = ("markup_tabla", "propuesta", "acta_apertura", "contrato", "otro")


# ── Text extraction ──────────────────────────────────────────────────────────

def _extract_xlsx(file_bytes: bytes) -> str:
    import openpyxl
    from io import BytesIO
    wb = openpyxl.load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
    lines = []
    for sheet in wb.worksheets:
        lines.append(f"=== Hoja: {sheet.title} ===")
        for row in sheet.iter_rows(values_only=True):
            parts = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if parts:
                lines.append(" | ".join(parts))
    return "\n".join(lines)


def _extract_docx(file_bytes: bytes) -> str:
    from docx import Document
    from io import BytesIO
    doc = Document(BytesIO(file_bytes))
    lines = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        import pypdf
        from io import BytesIO
        reader = pypdf.PdfReader(BytesIO(file_bytes))
        pages = []
        for page in reader.pages[:50]:
            t = page.extract_text() or ""
            if t.strip():
                pages.append(t.strip())
        return "\n\n".join(pages)
    except Exception as e:
        logger.warning(f"pypdf extraction failed: {e}")
        return ""


async def _extract_image_ocr(file_bytes: bytes, mime_type: str) -> str:
    """Use existing OCRService for scanned images/PDFs."""
    try:
        from services.ocr_service import get_ocr_service
        svc = get_ocr_service()
        result = await svc.process_image(file_bytes, mime_type)
        if result:
            parts = [result.get("objeto", ""), result.get("observaciones", "")]
            return " ".join(p for p in parts if p)
    except Exception as e:
        logger.warning(f"OCR extraction failed: {e}")
    return ""


async def _extract_text(file_bytes: bytes, filename: str, mime_type: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in ("xlsx", "xls"):
        return await asyncio.to_thread(_extract_xlsx, file_bytes)
    if ext in ("docx", "doc"):
        return await asyncio.to_thread(_extract_docx, file_bytes)
    if ext == "pdf" or mime_type == "application/pdf":
        text = await asyncio.to_thread(_extract_pdf, file_bytes)
        if len(text.strip()) < 100:
            # Likely scanned — fallback to OCR
            text = await _extract_image_ocr(file_bytes, mime_type)
        return text
    if mime_type.startswith("image/"):
        return await _extract_image_ocr(file_bytes, mime_type)
    return ""


# ── Chunking ─────────────────────────────────────────────────────────────────

def _chunk_text(text: str) -> list[str]:
    """Split text into overlapping chunks of ~CHUNK_SIZE chars."""
    text = text.strip()
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + CHUNK_SIZE
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = end - CHUNK_OVERLAP
    return chunks


# ── Embedding (Gemini text-embedding-004 via HTTP, 768-dim) ──────────────────

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")
EMBED_MODEL = "gemini-embedding-001"
EMBED_DIM = 768


async def _embed(text: str) -> Optional[list]:
    """Embed text using Gemini text-embedding-004. Returns 768-dim list or None."""
    if not GEMINI_API_KEY:
        logger.warning("GEMINI_API_KEY not set — embeddings disabled")
        return None
    import aiohttp
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{EMBED_MODEL}:embedContent?key={GEMINI_API_KEY}"
    payload = {
        "model": f"models/{EMBED_MODEL}",
        "content": {"parts": [{"text": text[:2000]}]},
        "taskType": "RETRIEVAL_DOCUMENT",
    }
    try:
        async with aiohttp.ClientSession() as session:
            async with session.post(url, json=payload, timeout=aiohttp.ClientTimeout(total=30)) as r:
                if r.status != 200:
                    logger.warning(f"Gemini embed error {r.status}: {await r.text()}")
                    return None
                data = await r.json()
                return data["embedding"]["values"]
    except Exception as e:
        logger.warning(f"Gemini embed failed: {e}")
        return None


async def _embed_query(text: str) -> Optional[list]:
    """Embed a search query (taskType RETRIEVAL_QUERY)."""
    if not GEMINI_API_KEY:
        return None
    import aiohttp
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{EMBED_MODEL}:embedContent?key={GEMINI_API_KEY}"
    payload = {
        "model": f"models/{EMBED_MODEL}",
        "content": {"parts": [{"text": text[:2000]}]},
        "taskType": "RETRIEVAL_QUERY",
    }
    try:
        async with aiohttp.ClientSession() as session:
            async with session.post(url, json=payload, timeout=aiohttp.ClientTimeout(total=30)) as r:
                if r.status != 200:
                    return None
                data = await r.json()
                return data["embedding"]["values"]
    except Exception as e:
        logger.warning(f"Gemini query embed failed: {e}")
        return None


# ── Service ──────────────────────────────────────────────────────────────────

class UMKnowledgeService:

    async def ensure_indexes(self, db: AsyncIOMotorDatabase):
        await db.um_knowledge.create_index("doc_id")
        await db.um_knowledge.create_index("tipo")
        await db.um_knowledge.create_index("created_at")

    async def ingest_file(
        self,
        db: AsyncIOMotorDatabase,
        file_bytes: bytes,
        filename: str,
        mime_type: str,
        tipo: str,
        metadata: Optional[dict] = None,
    ) -> dict:
        """Extract, chunk, embed and store a document. Returns summary."""
        if tipo not in TIPOS:
            tipo = "otro"

        text = await _extract_text(file_bytes, filename, mime_type)
        if not text.strip():
            return {"doc_id": None, "chunks_created": 0, "error": "No se pudo extraer texto del documento"}

        chunks = _chunk_text(text)
        if not chunks:
            return {"doc_id": None, "chunks_created": 0, "error": "Texto vacío tras chunking"}

        doc_id = str(uuid.uuid4())
        now = datetime.now(timezone.utc)
        docs = []

        for i, chunk in enumerate(chunks):
            vec = await _embed(chunk)
            docs.append({
                "doc_id": doc_id,
                "filename": filename,
                "tipo": tipo,
                "chunk_index": i,
                "chunk_text": chunk,
                "vector": vec,
                "metadata": metadata or {},
                "created_at": now,
            })

        if docs:
            await db.um_knowledge.insert_many(docs)

        logger.info(f"um_knowledge: ingested '{filename}' → {len(docs)} chunks (doc_id={doc_id})")
        return {"doc_id": doc_id, "chunks_created": len(docs), "tipo": tipo, "filename": filename}

    async def search(
        self,
        db: AsyncIOMotorDatabase,
        query: str,
        tipo: Optional[str] = None,
        top_k: int = 5,
    ) -> list[dict]:
        """Semantic search over um_knowledge chunks. Returns top-K with scores."""
        try:
            import numpy as np
        except ImportError:
            logger.warning("numpy not available — um_knowledge search disabled")
            return []

        vec = await _embed_query(query)
        if vec is None:
            # Fallback: simple text grep (no embeddings)
            filt = {"chunk_text": {"$regex": query[:50], "$options": "i"}}
            if tipo:
                filt["tipo"] = tipo
            docs = await db.um_knowledge.find(filt).limit(top_k).to_list(top_k)
            for d in docs:
                d["score"] = 0.5
                d["_id"] = str(d["_id"])
            return docs

        q_arr = np.array(vec, dtype=np.float32)
        norm_q = float(np.linalg.norm(q_arr))
        if norm_q == 0:
            return []

        filt = {}
        if tipo:
            filt["tipo"] = tipo

        all_chunks = await db.um_knowledge.find(
            filt, {"_id": 1, "doc_id": 1, "filename": 1, "tipo": 1, "chunk_index": 1, "chunk_text": 1, "vector": 1, "metadata": 1}
        ).to_list(None)

        scores = []
        for ch in all_chunks:
            v = ch.get("vector")
            if not v:
                continue
            v_arr = np.array(v, dtype=np.float32)
            norm_v = float(np.linalg.norm(v_arr))
            if norm_v == 0:
                continue
            sim = float(np.dot(q_arr, v_arr) / (norm_q * norm_v))
            scores.append((ch, sim))

        scores.sort(key=lambda x: x[1], reverse=True)
        results = []
        for ch, sim in scores[:top_k]:
            ch["score"] = round(sim, 3)
            ch["_id"] = str(ch["_id"])
            ch.pop("vector", None)
            results.append(ch)
        return results

    async def search_documents(
        self,
        db: AsyncIOMotorDatabase,
        query: str,
        top_k: int = 4,
    ) -> list[dict]:
        """Search and group by document. Returns top_k unique docs with synthesis."""
        raw = await self.search(db, query, top_k=min(top_k * 8, 40))

        doc_map: dict = {}
        for ch in raw:
            did = ch["doc_id"]
            if did not in doc_map:
                doc_map[did] = {
                    "doc_id": did,
                    "filename": ch["filename"],
                    "tipo": ch["tipo"],
                    "best_score": ch["score"],
                    "matching_chunks": [ch["chunk_text"]],
                    "created_at": ch.get("created_at"),
                }
            else:
                doc_map[did]["matching_chunks"].append(ch["chunk_text"])

        sorted_docs = sorted(doc_map.values(), key=lambda d: d["best_score"], reverse=True)[:top_k]

        for doc in sorted_docs:
            doc["key_data"] = _synthesize_chunks(doc["matching_chunks"])
            doc["excerpt"] = doc["matching_chunks"][0][:400]
            del doc["matching_chunks"]

        return sorted_docs

    async def list_docs(self, db: AsyncIOMotorDatabase) -> list[dict]:
        """List documents grouped by doc_id with chunk counts."""
        pipeline = [
            {"$group": {
                "_id": "$doc_id",
                "filename": {"$first": "$filename"},
                "tipo": {"$first": "$tipo"},
                "chunks": {"$sum": 1},
                "metadata": {"$first": "$metadata"},
                "created_at": {"$first": "$created_at"},
            }},
            {"$sort": {"created_at": -1}},
        ]
        docs = await db.um_knowledge.aggregate(pipeline).to_list(200)
        for d in docs:
            d["doc_id"] = d.pop("_id")
        return docs

    async def delete_doc(self, db: AsyncIOMotorDatabase, doc_id: str) -> int:
        result = await db.um_knowledge.delete_many({"doc_id": doc_id})
        return result.deleted_count

    async def get_stats(self, db: AsyncIOMotorDatabase) -> dict:
        total_chunks = await db.um_knowledge.count_documents({})
        pipeline = [{"$group": {"_id": "$tipo", "chunks": {"$sum": 1}, "docs": {"$addToSet": "$doc_id"}}}]
        by_tipo_raw = await db.um_knowledge.aggregate(pipeline).to_list(20)
        by_tipo = {r["_id"]: {"chunks": r["chunks"], "docs": len(r["docs"])} for r in by_tipo_raw}
        total_docs_pipeline = [{"$group": {"_id": "$doc_id"}}]
        total_docs = len(await db.um_knowledge.aggregate(total_docs_pipeline).to_list(None))
        return {"total_chunks": total_chunks, "total_docs": total_docs, "by_tipo": by_tipo}


# ── Document synthesis helpers ────────────────────────────────────────────────

def _try_parse_num(val: str) -> Optional[float]:
    s = val.strip().replace("\xa0", "").replace(",", ".")
    s = re.sub(r"[$\s\t]", "", s)
    try:
        return float(s) if s else None
    except ValueError:
        return None


def _synthesize_chunks(chunks: list) -> dict:
    """Extract margins (0<n≤1) and amounts (n>999) from pipe-delimited chunk rows."""
    margins = []
    amounts = []
    seen_labels: set = set()

    for chunk_text in chunks:
        for line in chunk_text.split("\n"):
            if "===" in line or line.startswith("HEADERS:"):
                continue
            cells = [c.strip() for c in line.split("|")]
            if len(cells) < 2:
                continue
            label = cells[0]
            if not label or not label[:1].isalpha() or label in seen_labels:
                continue
            seen_labels.add(label)
            for cell in cells[1:]:
                n = _try_parse_num(cell)
                if n is None:
                    continue
                if 0 < abs(n) <= 1:
                    margins.append({"label": label[:50], "pct": round(n, 4)})
                    break
                elif abs(n) > 999:
                    amounts.append({"label": label[:50], "amount": round(n, 2)})
                    break

    return {
        "margins": margins[:8],
        "amounts": sorted(amounts, key=lambda x: -x["amount"])[:5],
    }


_instance: Optional[UMKnowledgeService] = None


def get_um_knowledge_service() -> UMKnowledgeService:
    global _instance
    if _instance is None:
        _instance = UMKnowledgeService()
    return _instance
