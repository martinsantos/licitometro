"""
KnowledgeService — Base de conocimiento empresarial multi-fuente.

Nodos (knowledge_nodes): agrupadores temáticos (SOFTWARE, RED INCENDIO, REDES, etc.)
Docs  (knowledge_docs):  documentos de cualquier fuente asignados a un nodo

Fuentes soportadas: upload, licitacion, antecedente, url, paste
Tipos de doc: oferta, adjudicacion, especificacion, marketing, precio, cierre, referencia, otro

Reutiliza:
  - extract_text_from_pdf_bytes()  de enrichment/pdf_zip_enricher.py
  - GroqEnrichmentService.interpret_knowledge_doc() para summary + entities
  - openpyxl / python-docx para XLSX/DOCX
"""
from __future__ import annotations

import asyncio
import logging
import re
import uuid
from datetime import datetime, timezone
from typing import Optional

from motor.motor_asyncio import AsyncIOMotorDatabase

logger = logging.getLogger("knowledge_service")

DOC_TYPES = ("oferta", "adjudicacion", "especificacion", "marketing", "precio", "cierre", "referencia", "otro")
SOURCES   = ("upload", "licitacion", "antecedente", "url", "paste")


# ── Text extraction ──────────────────────────────────────────────────────────

def _extract_xlsx(file_bytes: bytes) -> str:
    """Extrae texto de XLSX con HEADERS inline en cada chunk de 30 filas."""
    import openpyxl
    from io import BytesIO
    wb = openpyxl.load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        # Usar primera fila no-vacía como headers
        header_idx = 0
        for i, row in enumerate(rows):
            if any(c is not None and str(c).strip() for c in row):
                header_idx = i
                break
        headers = " | ".join(str(c).strip() if c is not None else "" for c in rows[header_idx])
        data_rows = rows[header_idx + 1:]
        # Chunk de 30 filas, con headers repetidos
        chunk_size = 30
        for i in range(0, max(1, len(data_rows)), chunk_size):
            chunk_rows = data_rows[i:i + chunk_size]
            lines = [f"=== Hoja: {sheet.title} ===", f"HEADERS: {headers}"]
            for row in chunk_rows:
                cells = [str(c).strip() if c is not None else "" for c in row]
                if any(c for c in cells):
                    lines.append(" | ".join(cells))
            parts.append("\n".join(lines))
    return "\n\n".join(parts)


def _extract_docx(file_bytes: bytes) -> str:
    from docx import Document
    from io import BytesIO
    doc = Document(BytesIO(file_bytes))
    lines = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_csv(file_bytes: bytes) -> str:
    import csv
    from io import StringIO
    text = file_bytes.decode("utf-8", errors="replace")
    reader = csv.reader(StringIO(text))
    rows = list(reader)
    if not rows:
        return ""
    headers = " | ".join(rows[0])
    lines = [f"HEADERS: {headers}"]
    for row in rows[1:]:
        lines.append(" | ".join(row))
    return "\n".join(lines)


async def _extract_pdf(file_bytes: bytes) -> str:
    try:
        from services.enrichment.pdf_zip_enricher import extract_text_from_pdf_bytes
        return await asyncio.to_thread(extract_text_from_pdf_bytes, file_bytes)
    except Exception:
        pass
    try:
        import pypdf
        from io import BytesIO
        reader = pypdf.PdfReader(BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages[:50]]
        return "\n\n".join(p.strip() for p in pages if p.strip())
    except Exception as e:
        logger.warning(f"PDF extraction failed: {e}")
        return ""


async def extract_text(file_bytes: bytes, filename: str, mime_type: str = "") -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in ("xlsx", "xls"):
        return await asyncio.to_thread(_extract_xlsx, file_bytes)
    if ext in ("docx", "doc"):
        return await asyncio.to_thread(_extract_docx, file_bytes)
    if ext == "csv":
        return await asyncio.to_thread(_extract_csv, file_bytes)
    if ext == "pdf" or mime_type == "application/pdf":
        return await _extract_pdf(file_bytes)
    if mime_type.startswith("image/"):
        try:
            from services.ocr_service import get_ocr_service
            svc = get_ocr_service()
            result = await svc.process_image(file_bytes, mime_type)
            if result:
                return " ".join(v for v in result.values() if isinstance(v, str) and v)
        except Exception:
            pass
    return file_bytes.decode("utf-8", errors="replace")[:50_000]


async def _fetch_url_text(url: str) -> str:
    try:
        import aiohttp
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=aiohttp.ClientTimeout(total=30),
                                   ssl=False, allow_redirects=True) as r:
                content_type = r.headers.get("Content-Type", "")
                if "pdf" in content_type or url.lower().endswith(".pdf"):
                    data = await r.read()
                    return await _extract_pdf(data)
                html = await r.text(errors="replace")
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(html, "html.parser")
                for tag in soup(["script", "style", "nav", "header", "footer"]):
                    tag.decompose()
                return soup.get_text(" ", strip=True)[:20_000]
    except Exception as e:
        logger.warning(f"URL fetch failed {url}: {e}")
        return ""


# ── Slug generation ───────────────────────────────────────────────────────────

def _slugify(text: str) -> str:
    text = text.lower().strip()
    text = re.sub(r"[áàâä]", "a", text)
    text = re.sub(r"[éèêë]", "e", text)
    text = re.sub(r"[íìîï]", "i", text)
    text = re.sub(r"[óòôö]", "o", text)
    text = re.sub(r"[úùûü]", "u", text)
    text = re.sub(r"ñ", "n", text)
    text = re.sub(r"[^a-z0-9]+", "_", text)
    return text.strip("_")[:50]


# ── Main Service ──────────────────────────────────────────────────────────────

class KnowledgeService:

    async def ensure_indexes(self, db: AsyncIOMotorDatabase):
        await db.knowledge_nodes.create_index("slug", unique=True)
        await db.knowledge_nodes.create_index("created_at")
        await db.knowledge_docs.create_index("node_id")
        await db.knowledge_docs.create_index("doc_type")
        await db.knowledge_docs.create_index("source")
        await db.knowledge_docs.create_index("uploaded_at")
        await db.knowledge_docs.create_index([("full_text", "text")], default_language="none")

    # ── Node CRUD ─────────────────────────────────────────────────────────────

    async def create_node(self, db: AsyncIOMotorDatabase, name: str, color: str = "#6366f1",
                          description: str = "", keywords: list[str] = None) -> dict:
        slug = _slugify(name)
        # Ensure unique slug
        existing = await db.knowledge_nodes.find_one({"slug": slug})
        if existing:
            slug = f"{slug}_{uuid.uuid4().hex[:4]}"
        doc = {
            "name": name,
            "slug": slug,
            "color": color,
            "description": description,
            "keywords": keywords or [],
            "doc_count": 0,
            "chunk_count": 0,
            "created_at": datetime.now(timezone.utc),
        }
        result = await db.knowledge_nodes.insert_one(doc)
        doc["_id"] = str(result.inserted_id)
        return doc

    async def list_nodes(self, db: AsyncIOMotorDatabase) -> list[dict]:
        nodes = await db.knowledge_nodes.find().sort("name", 1).to_list(200)
        for n in nodes:
            n["_id"] = str(n["_id"])
        return nodes

    async def update_node(self, db: AsyncIOMotorDatabase, node_id: str, **fields) -> bool:
        from bson import ObjectId
        allowed = {"name", "color", "description", "keywords"}
        update = {k: v for k, v in fields.items() if k in allowed}
        if not update:
            return False
        result = await db.knowledge_nodes.update_one({"_id": ObjectId(node_id)}, {"$set": update})
        return result.modified_count > 0

    async def delete_node(self, db: AsyncIOMotorDatabase, node_id: str) -> int:
        from bson import ObjectId
        oid = ObjectId(node_id)
        docs_deleted = await db.knowledge_docs.delete_many({"node_id": oid})
        await db.knowledge_nodes.delete_one({"_id": oid})
        return docs_deleted.deleted_count

    async def _update_node_counts(self, db: AsyncIOMotorDatabase, node_id):
        from bson import ObjectId
        oid = ObjectId(str(node_id)) if not hasattr(node_id, "binary") else node_id
        doc_count = await db.knowledge_docs.count_documents({"node_id": oid})
        chunk_count_result = await db.knowledge_docs.aggregate([
            {"$match": {"node_id": oid}},
            {"$project": {"n": {"$size": {"$ifNull": ["$chunks", []]}}}},
            {"$group": {"_id": None, "total": {"$sum": "$n"}}},
        ]).to_list(1)
        chunk_count = chunk_count_result[0]["total"] if chunk_count_result else 0
        await db.knowledge_nodes.update_one(
            {"_id": oid}, {"$set": {"doc_count": doc_count, "chunk_count": chunk_count}}
        )

    # ── Ingestion ─────────────────────────────────────────────────────────────

    async def _store_doc(self, db: AsyncIOMotorDatabase, node_id, source: str,
                         doc_type: str, full_text: str, filename: str,
                         source_ref_id: str = "", url: str = "",
                         entities: dict = None, mime_type: str = "") -> dict:
        from bson import ObjectId
        from services.groq_enrichment import get_groq_enrichment_service

        if doc_type not in DOC_TYPES:
            doc_type = "otro"

        # AI interpretation
        groq = get_groq_enrichment_service()
        ai_entities = await groq.interpret_knowledge_doc(full_text, doc_type, filename)
        merged_entities = {**(entities or {}), **ai_entities} if ai_entities else (entities or {})
        summary = merged_entities.pop("summary", "") or ""

        # Chunk text for retrieval (preserve HEADERS blocks, ~1800 chars)
        chunks = self._chunk(full_text)

        oid = ObjectId(str(node_id)) if not hasattr(node_id, "binary") else node_id

        doc = {
            "node_id": oid,
            "source": source,
            "source_ref_id": source_ref_id,
            "doc_type": doc_type,
            "filename": filename,
            "mime_type": mime_type,
            "url": url,
            "summary": summary,
            "entities": merged_entities,
            "full_text": full_text[:100_000],
            "chunks": [{"chunk_index": i, "text": c, "embedding": None} for i, c in enumerate(chunks)],
            "uploaded_at": datetime.now(timezone.utc),
        }
        result = await db.knowledge_docs.insert_one(doc)
        doc["_id"] = str(result.inserted_id)
        doc["node_id"] = str(oid)
        await self._update_node_counts(db, oid)
        logger.info(f"knowledge: ingested '{filename}' ({doc_type}) → {len(chunks)} chunks, node={node_id}")
        return doc

    def _chunk(self, text: str) -> list[str]:
        """Chunk respetando bloques HEADERS (no parte filas de tabla a mitad)."""
        if not text.strip():
            return []
        # Si tiene bloques HEADERS, dividir por bloques (cada bloque es su propio chunk)
        if "HEADERS:" in text:
            blocks = re.split(r"\n\n(?===)", text)
            return [b.strip() for b in blocks if b.strip()]
        # Texto libre: chunking por caracteres con overlap
        CHUNK_SIZE, OVERLAP = 1800, 120
        chunks, start = [], 0
        while start < len(text):
            end = start + CHUNK_SIZE
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            if end >= len(text):
                break
            start = end - OVERLAP
        return chunks

    async def ingest_file(self, db: AsyncIOMotorDatabase, file_bytes: bytes,
                          filename: str, mime_type: str, node_id: str,
                          doc_type: str = "otro") -> dict:
        text = await extract_text(file_bytes, filename, mime_type)
        if not text.strip():
            return {"error": "No se pudo extraer texto del documento"}
        return await self._store_doc(db, node_id, "upload", doc_type, text, filename, mime_type=mime_type)

    async def ingest_url(self, db: AsyncIOMotorDatabase, url: str,
                         node_id: str, doc_type: str = "referencia") -> dict:
        text = await _fetch_url_text(url)
        if not text.strip():
            return {"error": f"No se pudo extraer texto de {url}"}
        filename = url.split("/")[-1][:80] or url[:80]
        return await self._store_doc(db, node_id, "url", doc_type, text, filename, url=url)

    async def ingest_paste(self, db: AsyncIOMotorDatabase, text: str, title: str,
                           node_id: str, doc_type: str = "otro") -> dict:
        if not text.strip():
            return {"error": "Texto vacío"}
        return await self._store_doc(db, node_id, "paste", doc_type, text, title or "Texto pegado")

    async def ingest_licitacion(self, db: AsyncIOMotorDatabase, lic_doc: dict,
                                node_id: str, doc_type: str = "oferta") -> dict:
        """Importa una licitación existente como doc de conocimiento (sin llamar a Groq)."""
        texto = " ".join(filter(None, [
            lic_doc.get("objeto", ""),
            lic_doc.get("title", ""),
            (lic_doc.get("description") or "")[:3000],
        ]))
        entities = {
            "clientes": [lic_doc.get("organization", "")],
            "monto_total": lic_doc.get("budget"),
            "fecha": str(lic_doc.get("publication_date", ""))[:7] or None,
            "resultado": lic_doc.get("workflow_state"),
        }
        filename = lic_doc.get("objeto") or lic_doc.get("title") or "Licitación"
        filename = filename[:80]
        return await self._store_doc(
            db, node_id, "licitacion", doc_type, texto, filename,
            source_ref_id=str(lic_doc.get("_id", "")), entities=entities,
        )

    # ── Docs ──────────────────────────────────────────────────────────────────

    async def list_docs(self, db: AsyncIOMotorDatabase, node_id: str,
                        doc_type: str = None, source: str = None,
                        limit: int = 50, skip: int = 0) -> list[dict]:
        from bson import ObjectId
        filt: dict = {"node_id": ObjectId(node_id)}
        if doc_type:
            filt["doc_type"] = doc_type
        if source:
            filt["source"] = source
        docs = await db.knowledge_docs.find(
            filt, {"full_text": 0, "chunks": 0}
        ).sort("uploaded_at", -1).skip(skip).limit(limit).to_list(limit)
        for d in docs:
            d["_id"] = str(d["_id"])
            d["node_id"] = str(d["node_id"])
        return docs

    async def delete_doc(self, db: AsyncIOMotorDatabase, doc_id: str) -> bool:
        from bson import ObjectId
        doc = await db.knowledge_docs.find_one({"_id": ObjectId(doc_id)}, {"node_id": 1})
        if not doc:
            return False
        await db.knowledge_docs.delete_one({"_id": ObjectId(doc_id)})
        await self._update_node_counts(db, doc["node_id"])
        return True

    # ── Search ────────────────────────────────────────────────────────────────

    async def search(self, db: AsyncIOMotorDatabase, q: str,
                     node_id: str = None, doc_type: str = None,
                     limit: int = 10) -> list[dict]:
        filt: dict = {}
        if node_id:
            from bson import ObjectId
            filt["node_id"] = ObjectId(node_id)
        if doc_type:
            filt["doc_type"] = doc_type

        results = []

        # $text search
        if q:
            text_filt = {**filt, "$text": {"$search": q}}
            docs = await db.knowledge_docs.find(
                text_filt,
                {"score": {"$meta": "textScore"}, "full_text": 0},
            ).sort([("score", {"$meta": "textScore"})]).limit(limit).to_list(limit)
            results = docs

        # Fallback: regex if no text results
        if not results:
            regex_filt = {**filt, "full_text": {"$regex": q[:50], "$options": "i"}}
            results = await db.knowledge_docs.find(
                regex_filt, {"full_text": 0}
            ).limit(limit).to_list(limit)

        for d in results:
            d["_id"] = str(d["_id"])
            d["node_id"] = str(d["node_id"])
            # Return relevant chunks only
            if "chunks" in d and q:
                q_lower = q.lower()
                relevant = [c for c in d["chunks"] if q_lower in c.get("text", "").lower()]
                d["chunks"] = relevant[:3] if relevant else d["chunks"][:2]
        return results

    async def get_context_for_cotizar(self, db: AsyncIOMotorDatabase,
                                      category: str = "", objeto: str = "") -> str:
        """Busca el nodo más relevante y devuelve string de contexto para inyectar al prompt."""
        nodes = await db.knowledge_nodes.find().to_list(100)
        if not nodes:
            return ""

        query_words = set(re.findall(r"\w{3,}", (category + " " + objeto).lower()))

        def overlap(node: dict) -> int:
            node_words = set(re.findall(r"\w{3,}", (
                " ".join(node.get("keywords", [])) + " " +
                node.get("name", "") + " " +
                node.get("description", "")
            ).lower()))
            return len(query_words & node_words)

        best = max(nodes, key=lambda n: (overlap(n), n.get("doc_count", 0)))
        if overlap(best) == 0:
            return ""

        from bson import ObjectId
        docs = await db.knowledge_docs.find(
            {"node_id": best["_id"], "doc_type": {"$in": ["precio", "adjudicacion", "oferta"]}},
            {"summary": 1, "entities": 1, "filename": 1, "doc_type": 1},
        ).sort("uploaded_at", -1).limit(5).to_list(5)

        if not docs:
            return ""

        lines = [f"CONOCIMIENTO INTERNO — Nodo: {best['name']} ({best.get('doc_count', 0)} docs)"]
        for d in docs:
            e = d.get("entities", {})
            line = f"  [{d['doc_type']}] {d['filename']}: {(d.get('summary') or '')[:100]}"
            if e.get("monto_total"):
                line += f" · ${e['monto_total']:,.0f}"
            if e.get("resultado"):
                line += f" · {e['resultado']}"
            lines.append(line)
        return "\n".join(lines)


_instance: Optional[KnowledgeService] = None


def get_knowledge_service() -> KnowledgeService:
    global _instance
    if _instance is None:
        _instance = KnowledgeService()
    return _instance
